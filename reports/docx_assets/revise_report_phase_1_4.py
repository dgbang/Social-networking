from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "bao_cao_phase_1_4_word_bao_cao_revised.docx"
ASSET_DIR = ROOT / "docx_assets"
FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")


PALETTE = {
    "ink": "#17324D",
    "muted": "#5D6B7A",
    "blue": "#2E74B5",
    "blue_soft": "#EAF2FA",
    "green": "#2E7D57",
    "green_soft": "#EAF6EF",
    "gold": "#A66A00",
    "gold_soft": "#FFF4DF",
    "gray": "#D8DEE6",
    "gray_soft": "#F5F7FA",
    "white": "#FFFFFF",
}


def font(size, bold=False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(path), size)


def hex_to_rgb(value):
    value = value.lstrip("#")
    return tuple(int(value[i:i + 2], 16) for i in (0, 2, 4))


def wrap_text(draw, text, font_obj, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def centered_text(draw, box, text, font_obj, fill=PALETTE["ink"], max_width=None, line_gap=8):
    x1, y1, x2, y2 = box
    max_width = max_width or (x2 - x1 - 28)
    lines = []
    for raw_line in text.split("\n"):
        lines.extend(wrap_text(draw, raw_line, font_obj, max_width) or [""])
    heights = [draw.textbbox((0, 0), line, font=font_obj)[3] - draw.textbbox((0, 0), line, font=font_obj)[1] for line in lines]
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font_obj)[2]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font_obj, fill=hex_to_rgb(fill))
        y += h + line_gap


def draw_box(draw, box, title, body=None, fill=PALETTE["gray_soft"], outline=PALETTE["blue"], radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=3)
    x1, y1, x2, y2 = box
    title_font = font(28, True)
    body_font = font(23)
    if body:
        draw.text((x1 + 24, y1 + 20), title, font=title_font, fill=hex_to_rgb(PALETTE["ink"]))
        y = y1 + 62
        for line in wrap_text(draw, body, body_font, x2 - x1 - 48):
            draw.text((x1 + 24, y), line, font=body_font, fill=hex_to_rgb(PALETTE["muted"]))
            y += 31
    else:
        centered_text(draw, box, title, title_font)


def arrow(draw, start, end, color=PALETTE["ink"], width=4):
    draw.line([start, end], fill=hex_to_rgb(color), width=width)
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 18
    p1 = (ex, ey)
    p2 = (ex - ux * size + px * size * 0.55, ey - uy * size + py * size * 0.55)
    p3 = (ex - ux * size - px * size * 0.55, ey - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=hex_to_rgb(color))


def label(draw, pos, text, size=22, color=PALETTE["muted"]):
    draw.text(pos, text, font=font(size), fill=hex_to_rgb(color))


def make_component_diagram(path):
    img = Image.new("RGB", (1800, 1080), hex_to_rgb(PALETTE["white"]))
    draw = ImageDraw.Draw(img)
    title = font(40, True)
    draw.text((60, 44), "Biểu đồ component tổng quát", font=title, fill=hex_to_rgb(PALETTE["ink"]))
    label(draw, (60, 98), "Các component được tách theo vai trò giao diện, API, nghiệp vụ, dữ liệu và dịch vụ ngoài.", 24)

    draw_box(draw, (110, 185, 490, 365), "React Client", "Vite, React Router, Redux Toolkit, Material UI, Axios", PALETTE["blue_soft"])
    draw_box(draw, (710, 165, 1090, 335), "Express API", "Routes, Controllers, Middlewares, Request validation", PALETTE["green_soft"], PALETTE["green"])
    draw_box(draw, (710, 415, 1090, 595), "Service Layer", "Auth, User, Friend, Post, Comment, Story services", PALETTE["gold_soft"], PALETTE["gold"])
    draw_box(draw, (710, 675, 1090, 835), "Data Access", "Sequelize models, migrations, associations", PALETTE["gray_soft"])
    draw_box(draw, (1285, 675, 1640, 835), "PostgreSQL", "Users, Sessions, Posts, Comments, Likes, Stories", PALETTE["blue_soft"])
    draw_box(draw, (1285, 205, 1640, 365), "External Services", "Cloudinary upload, SMTP mail, OAuth provider", PALETTE["gray_soft"], PALETTE["muted"])

    arrow(draw, (490, 275), (710, 250), PALETTE["blue"])
    label(draw, (540, 205), "REST API + JWT", 22)
    arrow(draw, (900, 335), (900, 415), PALETTE["green"])
    label(draw, (920, 360), "business call", 20)
    arrow(draw, (900, 595), (900, 675), PALETTE["gold"])
    label(draw, (920, 620), "ORM query", 20)
    arrow(draw, (1090, 755), (1285, 755), PALETTE["blue"])
    arrow(draw, (1090, 485), (1285, 285), PALETTE["muted"])
    label(draw, (1165, 390), "upload / email", 20)
    arrow(draw, (710, 515), (490, 315), PALETTE["muted"])
    label(draw, (520, 440), "JSON response", 20)

    legend_y = 930
    label(draw, (112, legend_y), "Ghi chú:", 24, PALETTE["ink"])
    label(draw, (218, legend_y), "Frontend chỉ gọi API; kiểm tra quyền, privacy và nghiệp vụ chính được xử lý ở backend.", 24)
    img.save(path, "PNG", quality=95)


def draw_actor(draw, x, y, name):
    pen = hex_to_rgb(PALETTE["ink"])
    draw.ellipse((x - 22, y, x + 22, y + 44), outline=pen, width=4)
    draw.line((x, y + 44, x, y + 118), fill=pen, width=4)
    draw.line((x - 48, y + 72, x + 48, y + 72), fill=pen, width=4)
    draw.line((x, y + 118, x - 42, y + 178), fill=pen, width=4)
    draw.line((x, y + 118, x + 42, y + 178), fill=pen, width=4)
    centered_text(draw, (x - 120, y + 195, x + 120, y + 260), name, font(24, True), max_width=230)


def use_case(draw, box, text, fill=PALETTE["blue_soft"]):
    draw.ellipse(box, fill=hex_to_rgb(fill), outline=hex_to_rgb(PALETTE["blue"]), width=3)
    centered_text(draw, box, text, font(24, True), max_width=250, line_gap=5)


def make_use_case_diagram(path):
    img = Image.new("RGB", (1800, 1120), hex_to_rgb(PALETTE["white"]))
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "Biểu đồ use case mức tổng quát", font=font(40, True), fill=hex_to_rgb(PALETTE["ink"]))
    label(draw, (60, 100), "Phạm vi thể hiện các chức năng đã hoàn thành từ Phase 1 đến Phase 4.", 24)

    system = (360, 165, 1440, 1010)
    draw.rounded_rectangle(system, radius=24, outline=hex_to_rgb(PALETTE["gray"]), width=4, fill=hex_to_rgb("#FBFCFE"))
    centered_text(draw, (620, 175, 1180, 230), "Social Networking Promax", font(30, True))

    draw_actor(draw, 180, 360, "Khách chưa đăng nhập")
    draw_actor(draw, 1620, 470, "Người dùng")

    cases = [
        ((455, 270, 735, 365), "Đăng ký\n tài khoản", PALETTE["green_soft"]),
        ((455, 410, 735, 505), "Đăng nhập /\n đăng xuất", PALETTE["green_soft"]),
        ((455, 550, 735, 645), "Quên và đặt lại\n mật khẩu", PALETTE["green_soft"]),
        ((880, 250, 1210, 345), "Cập nhật\n hồ sơ cá nhân", PALETTE["blue_soft"]),
        ((880, 370, 1210, 465), "Tìm kiếm và\n kết bạn", PALETTE["blue_soft"]),
        ((880, 490, 1210, 585), "Tạo và quản lý\n bài viết", PALETTE["blue_soft"]),
        ((880, 610, 1210, 705), "Xem bảng tin\n theo privacy", PALETTE["blue_soft"]),
        ((880, 730, 1210, 825), "Tương tác:\n like, comment, share", PALETTE["blue_soft"]),
        ((880, 850, 1210, 945), "Tạo và xem\n story 24 giờ", PALETTE["gold_soft"]),
    ]
    for box, text, fill in cases:
        use_case(draw, box, text, fill)

    def association(points):
        draw.line(points, fill=hex_to_rgb(PALETTE["muted"]), width=3, joint="curve")

    guest_anchor = (280, 460)
    for y in [318, 458, 598]:
        association([guest_anchor, (330, 460), (330, y), (455, y)])

    user_anchor = (1520, 570)
    for y in [298, 418, 538, 658, 778, 898]:
        association([user_anchor, (1480, 570), (1480, y), (1210, y)])

    label(draw, (385, 950), "Use case xác thực phục vụ cả khách và người dùng đã có tài khoản; các use case còn lại yêu cầu đăng nhập.", 22)
    img.save(path, "PNG", quality=95)


def sequence_text(draw, box, text, size=21):
    x1, y1, x2, y2 = box
    lines = []
    for line in text.split("\n"):
        lines.extend(wrap_text(draw, line, font(size), x2 - x1))
    y = y1
    for line in lines:
        draw.text((x1, y), line, font=font(size), fill=hex_to_rgb(PALETTE["ink"]))
        y += 27


def make_sequence_diagram(path):
    img = Image.new("RGB", (1900, 1120), hex_to_rgb(PALETTE["white"]))
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "Biểu đồ sequence: tạo bài viết có media", font=font(40, True), fill=hex_to_rgb(PALETTE["ink"]))
    label(draw, (60, 100), "Luồng tiêu biểu cho Phase 3, đồng thời thể hiện cách frontend, middleware, service, storage và database phối hợp.", 24)

    participants = [
        ("Người dùng", 170),
        ("React UI", 440),
        ("API Controller", 730),
        ("Auth Middleware", 1030),
        ("Post Service", 1325),
        ("Storage/DB", 1630),
    ]
    top = 185
    bottom = 1010
    for name, x in participants:
        draw.rounded_rectangle((x - 95, top, x + 95, top + 58), radius=14, fill=hex_to_rgb(PALETTE["blue_soft"]), outline=hex_to_rgb(PALETTE["blue"]), width=3)
        centered_text(draw, (x - 92, top + 5, x + 92, top + 55), name, font(22, True), max_width=170)
        draw.line((x, top + 58, x, bottom), fill=hex_to_rgb(PALETTE["gray"]), width=3)

    def seq_arrow(y, from_x, to_x, text, dashed=False, color=PALETTE["ink"]):
        if dashed:
            step = 18
            cur = from_x
            direction = 1 if to_x >= from_x else -1
            while (cur - to_x) * direction < 0:
                nxt = cur + direction * min(step, abs(to_x - cur))
                draw.line((cur, y, nxt, y), fill=hex_to_rgb(color), width=3)
                cur = nxt + direction * 10
            arrow(draw, (to_x - direction * 1, y), (to_x, y), color, 3)
        else:
            arrow(draw, (from_x, y), (to_x, y), color, 3)
        mid = min(from_x, to_x) + abs(to_x - from_x) / 2
        sequence_text(draw, (mid - 155, y - 47, mid + 155, y - 8), text, 19)

    steps = [
        (310, 170, 440, "1. Nhập nội dung,\nchọn ảnh/video"),
        (390, 440, 730, "2. POST /posts\nJWT + FormData"),
        (470, 730, 1030, "3. Xác thực token\nvà gắn req.user"),
        (550, 1030, 730, "4. Cho phép request", True, PALETTE["green"]),
        (630, 730, 1325, "5. Gọi createPost()"),
        (710, 1325, 1630, "6. Upload media,\nlưu Post/PostMedia"),
        (790, 1630, 1325, "7. Trả dữ liệu đã lưu", True, PALETTE["green"]),
        (870, 1325, 730, "8. Chuẩn hóa response", True, PALETTE["green"]),
        (950, 730, 440, "9. HTTP 201 + post", True, PALETTE["green"]),
        (1015, 440, 170, "10. Cập nhật feed\ntrên giao diện", True, PALETTE["green"]),
    ]
    for item in steps:
        if len(item) == 4:
            y, start, end, text = item
            seq_arrow(y, start, end, text)
        else:
            y, start, end, text, dashed, color = item
            seq_arrow(y, start, end, text, dashed, color)

    img.save(path, "PNG", quality=95)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def add_after(cursor, text="", style=None, bold=False, italic=False, color=None, size=None):
    paragraph = insert_paragraph_after(cursor, "", style)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        if size:
            run.font.size = Pt(size)
    if style and style.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_bullet_after(cursor, text):
    return add_after(cursor, text, "List Bullet")


def add_picture_after(cursor, image_path, width_inches=6.25):
    paragraph = insert_paragraph_after(cursor, "")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width_inches))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", image_path.stem.replace("_", " "))
    return paragraph


def caption_after(cursor, text):
    paragraph = add_after(cursor, text, "Caption", italic=True, color="#555555")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


def renumber_existing_sections(doc):
    for para in doc.paragraphs:
        text = para.text.strip()
        if para.style.name == "Heading 1":
            match = re.match(r"^(\d+)\.\s*(.+)$", text)
            if match and int(match.group(1)) >= 3:
                para.text = f"{int(match.group(1)) + 1}. {match.group(2)}"
        elif para.style.name == "Heading 2":
            match = re.match(r"^3\.(\d+)\.\s*(.+)$", text)
            if match:
                para.text = f"4.{match.group(1)}. {match.group(2)}"


def insert_overview_note(doc):
    marker = "Xây dựng chức năng story với cơ chế hết hạn sau 24 giờ."
    for para in doc.paragraphs:
        if para.text.strip() == marker:
            return add_bullet_after(
                para,
                "Bổ sung phần mô tả kiến trúc component, biểu đồ use case và biểu đồ sequence để làm rõ cách hệ thống được thiết kế."
            )
    return None


def insert_design_section(doc, component_img, use_case_img, sequence_img):
    anchor = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == "Heading 1" and para.text.strip().startswith("4. "):
            anchor = doc.paragraphs[i - 1]
            break
    if anchor is None:
        raise RuntimeError("Không tìm thấy vị trí chèn trước mục chức năng.")

    cursor = add_after(anchor, "3. Kiến trúc component và biểu đồ thiết kế", "Heading 1")
    cursor = add_after(
        cursor,
        "Phần này bổ sung mô tả kiến trúc và các biểu đồ phân tích theo góp ý của giảng viên. "
        "Các biểu đồ tập trung vào phạm vi đã hoàn thành đến hết Phase 4, gồm xác thực, hồ sơ, bạn bè, bài viết, tương tác và story."
    )

    cursor = add_after(cursor, "3.1. Component hệ thống", "Heading 2")
    cursor = add_after(
        cursor,
        "Hệ thống được tổ chức theo mô hình client-server. Frontend chịu trách nhiệm hiển thị giao diện và quản lý trạng thái người dùng, "
        "backend xử lý xác thực, phân quyền và nghiệp vụ, còn database lưu trữ dữ liệu quan hệ của mạng xã hội."
    )
    cursor = add_bullet_after(cursor, "Frontend: React/Vite, React Router, Redux Toolkit, Axios và Material UI để xây dựng các màn hình chính.")
    cursor = add_bullet_after(cursor, "Backend API: Express tổ chức theo route, controller, middleware và service để tách rõ luồng nhận request và xử lý nghiệp vụ.")
    cursor = add_bullet_after(cursor, "Service layer: gom các nghiệp vụ auth, user, friend, post, comment và story, giúp tránh đặt logic phức tạp trực tiếp trong controller.")
    cursor = add_bullet_after(cursor, "Data layer: Sequelize model, migration và association làm cầu nối đến PostgreSQL.")
    cursor = add_bullet_after(cursor, "Dịch vụ ngoài: Cloudinary phục vụ upload ảnh/video, SMTP phục vụ đặt lại mật khẩu và OAuth provider dùng cho đăng nhập Google nếu bật.")
    cursor = add_picture_after(cursor, component_img, 6.4)
    cursor = caption_after(cursor, "Hình 1. Biểu đồ component tổng quát của hệ thống.")

    cursor = add_after(cursor, "3.2. Biểu đồ use case", "Heading 2")
    cursor = add_after(
        cursor,
        "Biểu đồ use case mô tả các nhóm chức năng chính mà người dùng tương tác với hệ thống. "
        "Khách chưa đăng nhập chỉ dùng các chức năng xác thực, trong khi người dùng đã đăng nhập có thể quản lý hồ sơ, bạn bè, bài viết, tương tác và story."
    )
    cursor = add_picture_after(cursor, use_case_img, 6.4)
    cursor = caption_after(cursor, "Hình 2. Biểu đồ use case tổng quát cho các chức năng Phase 1 đến Phase 4.")

    cursor = add_after(cursor, "3.3. Biểu đồ sequence", "Heading 2")
    cursor = add_after(
        cursor,
        "Biểu đồ sequence dưới đây chọn luồng tạo bài viết có media làm ví dụ vì luồng này đi qua nhiều thành phần quan trọng: "
        "giao diện React, API controller, auth middleware, service nghiệp vụ, dịch vụ lưu trữ media và database."
    )
    cursor = add_picture_after(cursor, sequence_img, 6.4)
    cursor = caption_after(cursor, "Hình 3. Biểu đồ sequence cho luồng tạo bài viết có media.")


def tune_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for para in doc.paragraphs:
        if para.style.name == "Normal":
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.1
        elif para.style.name.startswith("Heading"):
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.keep_with_next = True


def main():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    component_img = ASSET_DIR / "component_diagram.png"
    use_case_img = ASSET_DIR / "use_case_diagram.png"
    sequence_img = ASSET_DIR / "sequence_diagram.png"

    make_component_diagram(component_img)
    make_use_case_diagram(use_case_img)
    make_sequence_diagram(sequence_img)

    doc = Document(DOCX_PATH)
    renumber_existing_sections(doc)
    insert_overview_note(doc)
    insert_design_section(doc, component_img, use_case_img, sequence_img)
    tune_document(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
