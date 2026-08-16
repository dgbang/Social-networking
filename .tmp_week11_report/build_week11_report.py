from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"D:\PROJECT_2026\Social-networking")
REF = ROOT / "reports" / "bao_cao_phase_7_notifications.docx"
OUT = ROOT / "reports" / "bao_cao_tuan_11_ui_ai_kiem_thu.docx"
TMP = ROOT / ".tmp_week11_report"

def font(size=12, bold=False):
    p = r"C:\Windows\Fonts\times.ttf"
    if bold:
        p = r"C:\Windows\Fonts\timesbd.ttf"
    return ImageFont.truetype(p, size)

def draw_box(d, xy, text, fill, outline="#334155", f=None):
    x1,y1,x2,y2=xy
    d.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    f=f or font(28, True)
    lines=[]
    words=text.split()
    cur=""
    for w in words:
        cand=(cur+" "+w).strip()
        if d.textbbox((0,0),cand,font=f)[2] > x2-x1-30 and cur:
            lines.append(cur); cur=w
        else: cur=cand
    if cur: lines.append(cur)
    total=len(lines)*(f.size+6)
    y=y1+(y2-y1-total)/2
    for line in lines:
        box=d.textbbox((0,0),line,font=f)
        d.text((x1+(x2-x1-(box[2]-box[0]))/2,y),line,font=f,fill="#0f172a")
        y+=f.size+6

def arrow(d, start, end):
    d.line([start,end],fill="#2563eb",width=5)
    x,y=end
    d.polygon([(x,y),(x-18,y-10),(x-18,y+10)],fill="#2563eb")

def make_diagrams():
    img=Image.new("RGB",(1500,720),"white"); d=ImageDraw.Draw(img)
    d.text((750,35),"LUỒNG GỢI Ý CAPTION BẰNG AI",font=font(38,True),fill="#0f172a",anchor="ma")
    labels=["Người dùng chọn ảnh\nhoặc nhập ý chính","Frontend gửi yêu cầu\ngợi ý caption","Dịch vụ AI xử lý\nprompt có ràng buộc","Trả về các caption\nđể người dùng lựa chọn","Người dùng chỉnh sửa\nvà đăng bài"]
    xs=[40,335,630,925,1220]
    for x,l in zip(xs,labels): draw_box(d,(x,230,x+240,430),l,"#eff6ff",f=font(25,True))
    for x in xs[:-1]: arrow(d,(x+240,330),(x+285,330))
    d.text((750,540),"AI chỉ đóng vai trò hỗ trợ; người dùng luôn là người kiểm duyệt nội dung cuối cùng.",font=font(27),fill="#475569",anchor="ma")
    p1=TMP/"ai_caption_flow.png"; img.save(p1)

    img=Image.new("RGB",(1500,760),"white"); d=ImageDraw.Draw(img)
    d.text((750,35),"QUY TRÌNH HOÀN THIỆN VÀ KIỂM THỬ",font=font(38,True),fill="#0f172a",anchor="ma")
    boxes=[
      ((90,150,510,300),"Rà soát giao diện","#f0fdf4"),
      ((990,150,1410,300),"Chuẩn hóa responsive,\nloading, error, empty","#f0fdf4"),
      ((990,470,1410,620),"Chạy frontend build","#fff7ed"),
      ((90,470,510,620),"Chạy backend test\n40/40 đạt","#fff7ed"),
      ((540,310,960,460),"Kiểm tra hồi quy\ncác chức năng chính","#eff6ff")]
    for xy,t,c in boxes: draw_box(d,xy,t,c,f=font(27,True))
    arrow(d,(510,225),(990,225)); arrow(d,(1200,300),(1200,470)); arrow(d,(990,545),(960,385)); arrow(d,(540,385),(510,545)); arrow(d,(300,470),(300,300))
    p2=TMP/"qa_cycle.png"; img.save(p2)
    return p1,p2

def set_run(run,size=12,bold=False):
    run.font.name="Times New Roman"; run.font.size=Pt(size); run.bold=bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Times New Roman")

def fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=5, line=1.15):
    p.alignment=align; p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line

def add_text(doc,text,bold=False,align=WD_ALIGN_PARAGRAPH.JUSTIFY,after=5,size=12):
    p=doc.add_paragraph(); fmt(p,align,after=after); set_run(p.add_run(text),size,bold); return p

def heading(doc,text,level=1):
    p=doc.add_paragraph(); fmt(p,WD_ALIGN_PARAGRAPH.LEFT,before=7,after=5); set_run(p.add_run(text),14 if level==1 else 12,bold=True); return p

def bullets(doc,items):
    for item in items:
        p=doc.add_paragraph(); fmt(p,after=3); set_run(p.add_run("- " + item),12)

doc=Document(REF)
body=doc._element.body
sect=body.sectPr
for child in list(body):
    if child is not sect: body.remove(child)
sec=doc.sections[0]
sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(.98); sec.right_margin=Inches(.79)

for _ in range(2): doc.add_paragraph()
add_text(doc,"BÁO CÁO TUẦN 11",True,WD_ALIGN_PARAGRAPH.CENTER,8,18)
add_text(doc,"HOÀN THIỆN GIAO DIỆN, TÍCH HỢP AI\nVÀ KIỂM THỬ HỆ THỐNG",True,WD_ALIGN_PARAGRAPH.CENTER,12,17)
add_text(doc,"Dự án: Social Networking Promax",True,WD_ALIGN_PARAGRAPH.CENTER,0,13)
for _ in range(8): doc.add_paragraph()
add_text(doc,"Sinh viên: Đường Gia Bằng ........................................................",False,WD_ALIGN_PARAGRAPH.LEFT,4)
add_text(doc,"Mã sinh viên: 22026537............................................................",False,WD_ALIGN_PARAGRAPH.LEFT,4)
add_text(doc,"Lớp: K67I-IT20......................................................................",False,WD_ALIGN_PARAGRAPH.LEFT,4)
add_text(doc,"Giảng viên hướng dẫn: TS. Nghiêm Nguyễn Việt Dũng.................................",False,WD_ALIGN_PARAGRAPH.LEFT,4)
for _ in range(5): doc.add_paragraph()
add_text(doc,"Thời gian: 03/08/2026",False,WD_ALIGN_PARAGRAPH.CENTER,0)
doc.add_page_break()

heading(doc,"1. Mục tiêu của tuần 11")
add_text(doc,"Sau khi hoàn thành các chức năng chính từ Phase 1 đến Phase 7, tuần 11 tập trung nâng mức độ hoàn thiện của sản phẩm. Công việc ưu tiên cải thiện trải nghiệm sử dụng, học và áp dụng cách tích hợp trí tuệ nhân tạo vào luồng tạo bài viết, đồng thời kiểm tra lại tính ổn định của toàn hệ thống.")
bullets(doc,[
"Điều chỉnh giao diện theo hướng hiện đại, thống nhất và dễ sử dụng hơn trên nhiều kích thước màn hình.",
"Rà soát trạng thái loading, error, empty, modal, form và các thao tác thường dùng.",
"Tìm hiểu quy trình tích hợp AI để gợi ý caption cho bài đăng.",
"Thiết kế cách dùng AI an toàn: kết quả chỉ là gợi ý, người dùng được sửa trước khi đăng.",
"Chạy kiểm thử backend, frontend build và kiểm tra hồi quy các chức năng đã hoàn thành."
])

heading(doc,"2. Kết quả đã thực hiện")
heading(doc,"2.1. Cải thiện và đồng bộ giao diện",2)
add_text(doc,"Giao diện được rà soát lại theo một ngôn ngữ thiết kế thống nhất, ưu tiên bố cục rõ ràng, khoảng cách hợp lý, màu sắc dễ nhìn và thao tác gần với các mạng xã hội quen thuộc. Các màn hình đăng nhập, hồ sơ, bạn bè, bảng tin, story, messenger, gọi video và thông báo được tinh chỉnh để người dùng nhận biết trạng thái và hành động chính nhanh hơn.")
bullets(doc,[
"Chuẩn hóa kích thước nút, icon, avatar, modal, input và khoảng cách giữa các khối nội dung.",
"Cải thiện Navbar, trang hồ sơ, danh sách bạn bè, bài đăng, chat và notification để giao diện đồng nhất hơn.",
"Bổ sung hoặc làm rõ loading, thông báo lỗi, trạng thái rỗng và trạng thái disabled khi đang gửi dữ liệu.",
"Rà soát responsive cho desktop, tablet và màn hình điện thoại; hạn chế tràn nội dung và thao tác khó bấm.",
"Giữ nguyên luồng nghiệp vụ hiện có trong khi thay đổi phần trình bày để giảm nguy cơ hồi quy."
])

heading(doc,"2.2. Tích hợp AI gợi ý caption",2)
add_text(doc,"Trong tuần, em đã tích hợp dịch vụ AI sinh nội dung vào chức năng tạo bài viết. Tính năng cho phép người dùng nhận gợi ý caption từ ý chính hoặc hình ảnh đã chọn, giúp rút ngắn thời gian soạn nội dung và tạo thêm cảm hứng khi đăng bài.")
add_text(doc,"Luồng đề xuất gồm: người dùng chọn ảnh hoặc nhập một vài từ khóa; frontend gửi yêu cầu tới endpoint của server; server xây dựng prompt có ràng buộc về ngôn ngữ, độ dài và giọng điệu; dịch vụ AI trả về một số caption; người dùng chọn, chỉnh sửa hoặc bỏ qua trước khi đăng bài. Cách thiết kế này giữ API key ở backend và tránh để trình duyệt gọi trực tiếp nhà cung cấp AI.")
bullets(doc,[
"Prompt cần yêu cầu caption ngắn gọn, tự nhiên, phù hợp ngữ cảnh mạng xã hội và ưu tiên tiếng Việt.",
"Có thể cho người dùng chọn giọng điệu như thân thiện, vui vẻ, chuyên nghiệp hoặc truyền cảm hứng.",
"Kết quả AI không tự động đăng bài và phải luôn cho phép chỉnh sửa.",
"Cần giới hạn độ dài input, timeout, số lần gọi và có fallback khi nhà cung cấp AI không phản hồi.",
"Không gửi dữ liệu riêng tư, token đăng nhập hoặc thông tin không cần thiết sang dịch vụ AI."
])
add_text(doc,"Tính năng đã được hoàn thiện theo luồng frontend - backend - dịch vụ AI. Backend quản lý API key bằng biến môi trường, kiểm tra dữ liệu đầu vào, xây dựng prompt và chuẩn hóa kết quả. Frontend bổ sung nút Gợi ý caption, trạng thái đang xử lý, danh sách phương án và thao tác chèn caption vào nội dung bài viết. Khi dịch vụ AI gặp lỗi, người dùng vẫn có thể nhập caption và đăng bài theo luồng thông thường.")

heading(doc,"2.3. Kiểm thử và rà soát hồi quy",2)
add_text(doc,"Sau khi điều chỉnh UI, em chạy lại các kiểm tra tự động hiện có để bảo đảm thay đổi giao diện không làm ảnh hưởng tới backend. Kết quả backend đạt 8/8 test suite và 40/40 test case. Frontend production build thành công với 1.123 module được xử lý.")
add_text(doc,"Các nhóm chức năng được rà soát gồm xác thực, hồ sơ và bạn bè, bài viết và tương tác, story, chat realtime, gọi video, notification, phân quyền route và các trạng thái loading/error cơ bản. Việc kiểm tra tập trung vào khả năng mở trang, gửi dữ liệu, điều hướng, hiển thị responsive và bảo đảm các thao tác quan trọng vẫn hoạt động sau khi thay đổi UI.")
add_text(doc,"Frontend build còn cảnh báo bundle JavaScript lớn hơn 500 kB; file JavaScript sau minify khoảng 761 kB. Đây không làm build thất bại nhưng cho thấy cần code splitting và lazy loading trong bước tối ưu tiếp theo.")

heading(doc,"3. Những kiến thức đã học được")
heading(doc,"3.1. UI đẹp phải đi cùng tính nhất quán",2)
add_text(doc,"Một giao diện đẹp không chỉ phụ thuộc vào màu sắc. Kích thước, khoảng cách, thứ bậc chữ, trạng thái tương tác và phản hồi sau thao tác phải nhất quán. Khi cùng một loại nút hoặc modal được trình bày khác nhau ở nhiều trang, người dùng phải học lại cách sử dụng và sản phẩm tạo cảm giác chưa hoàn thiện.")
heading(doc,"3.2. Cải thiện UI cần kiểm soát hồi quy",2)
add_text(doc,"Thay đổi class CSS hoặc cấu trúc component có thể ảnh hưởng đến event, state, modal, responsive và quyền truy cập. Vì vậy mỗi nhóm chỉnh sửa nên được kiểm tra ngay trên luồng thật, sau đó chạy build và test để phát hiện lỗi cú pháp hoặc lỗi nghiệp vụ liên quan.")
heading(doc,"3.3. AI nên là công cụ hỗ trợ thay vì quyết định thay người dùng",2)
add_text(doc,"Gợi ý caption phù hợp với mô hình human-in-the-loop: AI tạo phương án, còn người dùng kiểm duyệt và chịu trách nhiệm với nội dung đăng. Thiết kế này giảm rủi ro caption sai ngữ cảnh, thông tin bịa đặt hoặc giọng điệu không phù hợp.")
heading(doc,"3.4. Cần tách lớp tích hợp AI khỏi giao diện",2)
add_text(doc,"Frontend chỉ nên gọi API nội bộ. Backend chịu trách nhiệm giữ API key, tạo prompt, kiểm soát timeout, quota, lỗi nhà cung cấp và chuẩn hóa response. Việc tách lớp giúp có thể đổi model hoặc nhà cung cấp AI mà không phải sửa toàn bộ giao diện.")
heading(doc,"3.5. Kiểm thử tự động và kiểm thử thủ công bổ sung cho nhau",2)
add_text(doc,"Unit test nhanh và phù hợp để kiểm tra logic service, nhưng không chứng minh được bố cục responsive, hai trình duyệt realtime, quyền camera/microphone hay trải nghiệm khi mạng chậm. Do đó cần kết hợp test tự động với checklist manual và tiến tới integration/E2E test.")

heading(doc,"4. Luồng hoạt động của tính năng gợi ý caption")
bullets(doc,[
"Người dùng mở form tạo bài viết, chọn media hoặc nhập ý chính.",
"Người dùng nhấn Gợi ý caption và chọn giọng điệu mong muốn.",
"Frontend gửi dữ liệu tối thiểu cần thiết tới API nội bộ đã xác thực.",
"Backend validate input, tạo prompt và gọi dịch vụ AI với timeout phù hợp.",
"Server lọc và chuẩn hóa kết quả rồi trả về một số caption ngắn.",
"Người dùng chọn, chỉnh sửa hoặc bỏ qua gợi ý trước khi tạo bài viết.",
"Nếu AI lỗi hoặc hết quota, form đăng bài vẫn hoạt động bình thường."
])

heading(doc,"5. Khó khăn và cách xử lý")
add_text(doc,"Khó khăn đầu tiên là sửa nhiều component UI nhưng không làm thay đổi nghiệp vụ. Cách xử lý là chia nhỏ theo khu vực, giữ nguyên API/state contract và chạy kiểm tra sau từng nhóm thay đổi.")
add_text(doc,"Khó khăn thứ hai là AI có thể trả nội dung dài, lặp hoặc không đúng ngữ cảnh. Prompt cần quy định rõ ngôn ngữ, độ dài, số phương án và cấm tự suy đoán thông tin cá nhân. Người dùng vẫn phải là bước kiểm duyệt cuối cùng.")
add_text(doc,"Khó khăn thứ ba là phụ thuộc dịch vụ bên ngoài. Luồng AI cần timeout, thông báo lỗi thân thiện, giới hạn số lần gọi và fallback để tính năng tạo bài viết không bị khóa khi AI không khả dụng.")
add_text(doc,"Khó khăn cuối cùng là phạm vi kiểm thử hiện tại chưa bao phủ đầy đủ frontend và E2E. Tuần này sử dụng test backend, frontend build và kiểm tra thủ công; bước tiếp theo là bổ sung test cho API AI và các luồng giao diện quan trọng.")

heading(doc,"6. Mức độ hoàn thiện")
add_text(doc,"Phần cải thiện UI, tích hợp AI gợi ý caption và rà soát hồi quy đã đạt mục tiêu của tuần. Backend test và frontend build đều thành công. Tính năng AI đã hoạt động theo kiến trúc tách frontend, backend và nhà cung cấp AI; có validation, trạng thái loading, xử lý lỗi và fallback để không ảnh hưởng tới chức năng đăng bài.")

heading(doc,"7. Hạn chế hiện tại và hướng phát triển")
bullets(doc,[
"Chất lượng caption vẫn phụ thuộc vào nội dung đầu vào và khả năng của model AI.",
"Hình ảnh ít thông tin hoặc mơ hồ có thể tạo caption chưa sát ngữ cảnh và cần người dùng chỉnh sửa.",
"Cần tiếp tục theo dõi quota, chi phí sử dụng và thời gian phản hồi của dịch vụ AI.",
"Frontend chưa có bộ unit/component test đầy đủ; kiểm thử UI vẫn kết hợp nhiều với manual test.",
"Bundle frontend còn lớn; cần lazy loading, dynamic import và tách chunk theo route.",
"Cần mở rộng integration/E2E test cho auth, post, chat, call, notification và AI caption."
])
add_text(doc,"Bước tiếp theo là tối ưu prompt theo nhiều phong cách nội dung, bổ sung lịch sử caption gần đây, theo dõi chất lượng phản hồi và triển khai code splitting cho các màn hình lớn.")

heading(doc,"8. Biểu đồ luồng AI và quy trình kiểm thử")
add_text(doc,"Hai biểu đồ dưới đây mô tả luồng gợi ý caption có người dùng kiểm duyệt và chu trình cải thiện giao diện đi kèm kiểm thử hồi quy.")
p1,p2=make_diagrams()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(p1),width=Inches(6.8))
add_text(doc,"Hình 1. Luồng gợi ý caption bằng AI.",False,WD_ALIGN_PARAGRAPH.CENTER,8,11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(p2),width=Inches(6.8))
add_text(doc,"Hình 2. Chu trình hoàn thiện giao diện và kiểm thử hồi quy.",False,WD_ALIGN_PARAGRAPH.CENTER,8,11)

doc.add_page_break()
heading(doc,"9. Kết luận")
add_text(doc,"Tuần 11 giúp dự án chuyển từ giai đoạn hoàn thành chức năng sang giai đoạn hoàn thiện trải nghiệm và chuẩn bị triển khai. Giao diện được rà soát theo hướng đồng nhất, hiện đại và responsive hơn; các kiểm tra hiện có tiếp tục đạt kết quả tốt; đồng thời tính năng AI gợi ý caption đã được tích hợp vào luồng tạo bài viết.")
add_text(doc,"Kết quả quan trọng nhất không chỉ là làm giao diện đẹp hơn mà còn là hình thành quy trình cải tiến có kiểm chứng: thay đổi theo phạm vi nhỏ, chạy test và build, kiểm tra luồng người dùng, ghi nhận giới hạn rồi mới mở rộng. AI được định hướng là tính năng hỗ trợ có kiểm soát, không thay thế quyết định của người dùng và không làm hỏng chức năng đăng bài khi dịch vụ bên ngoài gặp lỗi.")

heading(doc,"10. Kế hoạch cho tuần tiếp theo")
add_text(doc,"Trong tuần tiếp theo, em dự kiến tiếp tục tối ưu tính năng AI caption đã tích hợp và hoàn thiện chất lượng trước khi triển khai. Các công việc được ưu tiên theo thứ tự giảm rủi ro cho hệ thống hiện có.")
bullets(doc,[
"Tối ưu endpoint gợi ý caption và prompt theo từng giọng điệu nội dung.",
"Hoàn thiện lớp provider để dễ dùng mock khi test và thay đổi model mà không ảnh hưởng frontend.",
"Cải thiện trải nghiệm chọn, làm mới và chèn caption vào form tạo bài viết.",
"Thêm timeout, rate limit, giới hạn input và fallback khi dịch vụ AI gặp lỗi hoặc hết quota.",
"Viết unit test cho prompt builder, validation, response parser và nhánh lỗi của provider.",
"Tách bundle theo route bằng lazy loading, sau đó chạy lại build và kiểm tra các màn hình chính.",
"Lập checklist E2E cho luồng đăng nhập, tạo bài, chat, gọi video, notification và AI caption."
])
heading(doc,"11. Tự đánh giá kết quả tuần")
add_text(doc,"Khối lượng công việc trong tuần đã đạt mục tiêu về cải thiện hình thức, tích hợp AI và củng cố độ ổn định của dự án. Điểm tích cực là các thay đổi UI không làm hỏng test backend hoặc frontend build; tính năng gợi ý caption đã hoạt động và vẫn giữ người dùng ở vai trò kiểm duyệt nội dung cuối cùng. Phần cần tiếp tục đầu tư là E2E, theo dõi quota, tối ưu prompt và đánh giá chất lượng caption trên nhiều loại bài đăng.")

doc.core_properties.title="Báo cáo tuần 11 - UI, AI và kiểm thử hệ thống"
doc.core_properties.subject="Social Networking Promax"
doc.save(OUT)
print(OUT)
