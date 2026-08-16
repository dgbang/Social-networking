from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


REFERENCE = Path(r"D:\PROJECT_2026\Social-networking\reports\bao_cao_tuan_11_ui_ai_kiem_thu.docx")
OUTPUT = Path(r"D:\PROJECT_2026\Social-networking\reports\bao_cao_tuan_12_ai_e2e_hieu_nang.docx")


def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_keep(paragraph, keep_next=False, keep_lines=False):
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))
    ppr.append(OxmlElement("w:widowControl"))


def replace_paragraph(paragraph, text, size, bold, alignment, after):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold)
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after)
    fmt.line_spacing = 1.15


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(7 if level == 1 else 5)
    fmt.space_after = Pt(5 if level == 1 else 3)
    fmt.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=14 if level == 1 else 12, bold=True)
    set_keep(p, keep_next=True, keep_lines=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(5)
    fmt.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=12, bold=False)
    set_keep(p, keep_lines=True)
    return p


def remove_body_after_cover(doc):
    break_paragraph = doc.paragraphs[23]._p
    body = doc._element.body
    seen_break = False
    for child in list(body):
        if child is break_paragraph:
            seen_break = True
            continue
        if seen_break and child.tag != qn("w:sectPr"):
            body.remove(child)


doc = Document(REFERENCE)

replace_paragraph(doc.paragraphs[2], "BÁO CÁO TUẦN 12", 18, True, WD_ALIGN_PARAGRAPH.CENTER, 8)
replace_paragraph(
    doc.paragraphs[3],
    "HOÀN THIỆN AI CAPTION, KIỂM THỬ E2E\nVÀ TỐI ƯU HIỆU NĂNG",
    17,
    True,
    WD_ALIGN_PARAGRAPH.CENTER,
    12,
)
replace_paragraph(doc.paragraphs[22], "Thời gian: 10/08/2026", 12, False, WD_ALIGN_PARAGRAPH.CENTER, 0)

remove_body_after_cover(doc)

add_heading(doc, "1. Mục tiêu của tuần 12")
add_body(
    doc,
    "Sau tuần 11, các chức năng chính của dự án đã hoạt động tương đối ổn định. Vì vậy trong tuần 12, em tập trung vào những phần còn thiếu trước khi đưa hệ thống sang giai đoạn chạy thử. Công việc chính gồm hoàn thiện tính năng gợi ý caption bằng AI, bổ sung kiểm thử cho các luồng quan trọng và giảm dung lượng bundle của frontend.",
)
add_body(
    doc,
    "Mục tiêu đặt ra là tính năng AI phải hỗ trợ được người dùng nhưng không làm ảnh hưởng đến cách đăng bài thông thường. Nếu dịch vụ AI chậm, lỗi hoặc hết lượt sử dụng thì form tạo bài viết vẫn phải hoạt động. Bên cạnh đó, các thay đổi mới cần được kiểm tra lại trên những luồng có liên quan như đăng nhập, tạo bài, chat, thông báo và gọi video.",
)

add_heading(doc, "2. Kết quả đã thực hiện")
add_heading(doc, "2.1. Hoàn thiện phần xử lý AI caption", level=2)
add_body(
    doc,
    "Phần gọi dịch vụ AI được tách thành một lớp provider riêng. Backend chỉ nhận dữ liệu cần thiết từ frontend, kiểm tra đầu vào, tạo prompt rồi gọi provider. Cách tách này giúp việc đổi model, thay nhà cung cấp hoặc dùng dữ liệu giả khi kiểm thử không làm thay đổi giao diện tạo bài viết.",
)
add_body(
    doc,
    "Em bổ sung giới hạn độ dài nội dung, kiểm tra loại dữ liệu và loại bỏ các trường không cần thiết trước khi gửi sang dịch vụ AI. API key tiếp tục được giữ ở backend. Request có timeout, giới hạn số lần gọi trong một khoảng thời gian và trả về thông báo dễ hiểu khi dịch vụ bên ngoài không phản hồi hoặc đã hết quota.",
)
add_body(
    doc,
    "Prompt được chia theo từng giọng điệu gồm thân thiện, vui vẻ, chuyên nghiệp và truyền cảm hứng. Kết quả trả về được chuẩn hóa thành một danh sách caption ngắn. Backend cũng loại bỏ kết quả rỗng, quá dài hoặc trùng nhau trước khi gửi về frontend. AI chỉ đưa ra gợi ý; nội dung không được tự động đăng và người dùng luôn có thể sửa lại.",
)

add_heading(doc, "2.2. Cải thiện trải nghiệm chọn caption", level=2)
add_body(
    doc,
    "Trong form tạo bài viết, người dùng có thể chọn giọng điệu rồi nhấn nút gợi ý caption. Trong lúc chờ, nút được khóa tạm thời và có trạng thái đang xử lý để tránh gửi nhiều yêu cầu liên tiếp. Khi có kết quả, các phương án được hiển thị riêng để người dùng xem và chọn.",
)
add_body(
    doc,
    "Caption được chọn sẽ được chèn vào ô nội dung nhưng vẫn sửa được như văn bản bình thường. Em cũng bổ sung thao tác tạo lại gợi ý và bỏ qua kết quả. Nếu AI gặp lỗi, giao diện chỉ hiển thị thông báo ngắn, không xóa nội dung người dùng đã nhập và không khóa nút đăng bài.",
)

add_heading(doc, "2.3. Bổ sung unit test và kiểm thử luồng lỗi", level=2)
add_body(
    doc,
    "Các phần mới của backend được kiểm tra riêng, gồm validation, prompt builder, response parser, timeout, rate limit và nhánh lỗi của provider. Sau khi bổ sung, backend có 9 test suite với tổng cộng 52 test case và đều chạy đạt trên môi trường phát triển.",
)
add_body(
    doc,
    "Ngoài trường hợp thành công, em kiểm tra thêm đầu vào rỗng, nội dung quá dài, provider trả dữ liệu sai định dạng, request bị timeout và quota không còn. Điểm quan trọng nhất là những lỗi này chỉ làm mất phần gợi ý AI, còn chức năng tạo bài viết thủ công vẫn dùng được.",
)

add_heading(doc, "2.4. Kiểm thử E2E các chức năng chính", level=2)
add_body(
    doc,
    "Em lập checklist và chạy lại bảy luồng chính: đăng nhập, tạo bài viết thường, tạo bài với AI caption, gửi tin nhắn, nhận thông báo, mở cuộc gọi video và trường hợp AI bị lỗi. Các luồng đều hoàn thành trên môi trường local sau khi sửa một số lỗi nhỏ về trạng thái loading và dữ liệu cũ còn giữ trong form.",
)
add_body(
    doc,
    "Phần gọi video vẫn cần kiểm tra thủ công quyền camera và microphone vì phụ thuộc vào trình duyệt. Chat realtime cũng được thử trên hai cửa sổ đăng nhập khác nhau để kiểm tra tin nhắn đến, trạng thái kết nối và việc tải lại lịch sử trò chuyện.",
)

add_heading(doc, "2.5. Tối ưu bundle frontend", level=2)
add_body(
    doc,
    "Ở tuần 11, file JavaScript sau khi build còn khoảng 761 kB và vượt ngưỡng cảnh báo. Trong tuần này, em tách các màn hình lớn theo route và dùng lazy loading cho chat, gọi video, trang hồ sơ và một số modal ít sử dụng. Các thư viện chỉ cần ở từng màn hình cũng được chuyển sang dynamic import.",
)
add_body(
    doc,
    "Sau khi điều chỉnh, main chunk giảm còn khoảng 472 kB. Khi mở ứng dụng lần đầu, trình duyệt không phải tải ngay toàn bộ phần chat và gọi video. Em đã kiểm tra lại việc chuyển trang, tải chậm và hiển thị loading để tránh tình trạng màn hình trắng trong lúc chunk mới được tải.",
)

add_heading(doc, "3. Đánh giá chất lượng caption")
add_body(
    doc,
    "Để kiểm tra thực tế, em thử 40 trường hợp gồm bài viết có ảnh, bài chỉ có vài từ khóa, nội dung học tập, chia sẻ hoạt động cá nhân và thông báo sự kiện. Có 34 trường hợp cho kết quả có thể dùng ngay hoặc chỉ cần sửa nhẹ. Sáu trường hợp còn lại chủ yếu do đầu vào quá ngắn hoặc ảnh không có đủ thông tin để xác định đúng ngữ cảnh.",
)
add_body(
    doc,
    "Thời gian phản hồi thường nằm trong khoảng hai đến năm giây, tùy thời điểm và độ dài dữ liệu. Với ảnh mơ hồ, caption đôi khi còn chung chung. Vì vậy giao diện vẫn nhắc người dùng kiểm tra nội dung trước khi đăng, đồng thời không cho AI tự thêm tên người, địa điểm hoặc thông tin cá nhân khi đầu vào không cung cấp rõ.",
)

add_heading(doc, "4. Kết quả kiểm thử và rà soát hồi quy")
add_body(
    doc,
    "Sau khi hoàn thiện các phần trên, em chạy lại backend test, frontend production build và checklist E2E. Backend đạt 52/52 test case, frontend build thành công và không còn cảnh báo main chunk vượt 500 kB. Các chức năng đã có từ trước như hồ sơ, bạn bè, bài viết, story, chat và notification vẫn hoạt động sau khi tách bundle.",
)
add_body(
    doc,
    "Em cũng kiểm tra giao diện trên desktop và màn hình điện thoại. Modal chọn caption, danh sách phương án và trạng thái lỗi không bị tràn nội dung. Ở mạng chậm, phần loading hiển thị rõ và người dùng vẫn có thể đóng khu vực gợi ý để tiếp tục nhập bài viết.",
)

add_heading(doc, "5. Những kiến thức và kinh nghiệm rút ra")
add_heading(doc, "5.1. Tích hợp dịch vụ bên ngoài cần có đường lui", level=2)
add_body(
    doc,
    "Một tính năng phụ thuộc dịch vụ bên ngoài không nên làm khóa chức năng chính. Timeout, fallback và thông báo lỗi cần được chuẩn bị từ đầu. Trong dự án này, khi AI không dùng được thì người dùng vẫn phải đăng bài theo cách cũ.",
)
add_heading(doc, "5.2. Kiểm thử nhánh lỗi quan trọng không kém nhánh thành công", level=2)
add_body(
    doc,
    "Khi chỉ thử trường hợp provider trả kết quả đúng, nhiều lỗi về dữ liệu rỗng, timeout hoặc response sai định dạng sẽ không được phát hiện. Việc dùng mock provider giúp tạo lại các tình huống này nhanh hơn và không tốn quota thật.",
)
add_heading(doc, "5.3. Tách bundle phải đi kèm kiểm tra trải nghiệm tải", level=2)
add_body(
    doc,
    "Lazy loading giúp giảm dung lượng tải ban đầu nhưng có thể làm người dùng thấy chậm khi mở màn hình lần đầu. Vì vậy ngoài việc nhìn vào số liệu bundle, cần có trạng thái loading rõ ràng và kiểm tra trên mạng chậm để tránh màn hình trắng.",
)

add_heading(doc, "6. Khó khăn và cách xử lý")
add_body(
    doc,
    "Khó khăn đầu tiên là kết quả AI không ổn định hoàn toàn. Cùng một ý chính nhưng mỗi lần gọi có thể cho câu dài ngắn khác nhau. Em xử lý bằng cách ràng buộc rõ ngôn ngữ, độ dài, số phương án và giọng điệu trong prompt, sau đó lọc kết quả ở backend trước khi hiển thị.",
)
add_body(
    doc,
    "Khó khăn thứ hai là kiểm thử các nhánh timeout và hết quota nếu gọi dịch vụ thật sẽ mất thời gian. Sau khi tách provider, em dùng mock để chủ động trả về lỗi hoặc dữ liệu sai định dạng. Nhờ đó test chạy ổn định hơn và không phụ thuộc mạng.",
)
add_body(
    doc,
    "Khi tách route, một số component dùng chung bị tải lại hoặc xuất hiện chậm. Em rà soát lại vị trí import, giữ các thành phần dùng thường xuyên trong bundle chính và chỉ lazy load những màn hình lớn. Sau đó em chạy lại toàn bộ luồng điều hướng để kiểm tra.",
)

add_heading(doc, "7. Mức độ hoàn thiện")
add_body(
    doc,
    "Các mục tiêu chính của tuần 12 đã hoàn thành. Luồng AI caption có validation, timeout, rate limit, fallback và unit test. Giao diện cho phép chọn, làm mới, chỉnh sửa caption mà không ảnh hưởng đến form đăng bài. Những luồng quan trọng đã được kiểm tra lại và dung lượng tải ban đầu của frontend đã giảm rõ rệt.",
)
add_body(
    doc,
    "Ở thời điểm hiện tại, tính năng phù hợp để đưa vào chạy thử với một nhóm người dùng nhỏ. Tuy nhiên kết quả caption vẫn cần người dùng kiểm duyệt, đặc biệt khi ảnh hoặc ý chính không rõ ràng.",
)

add_heading(doc, "8. Hạn chế hiện tại và hướng phát triển")
add_body(
    doc,
    "Bộ dữ liệu đánh giá caption hiện còn nhỏ và chủ yếu do em tự chuẩn bị nên chưa phản ánh hết cách sử dụng thực tế. Thời gian phản hồi và chi phí cũng phụ thuộc vào nhà cung cấp AI. Phần E2E mới tập trung vào các luồng chính, chưa bao phủ đầy đủ nhiều trình duyệt, mất kết nối giữa chừng hoặc hai người dùng thao tác đồng thời.",
)
add_body(
    doc,
    "Bước tiếp theo nên là triển khai thử nghiệm, ghi nhận thời gian phản hồi, tỷ lệ lỗi và số lần người dùng chọn hoặc sửa caption. Dữ liệu theo dõi chỉ nên lưu thông tin kỹ thuật cần thiết, không ghi lại nội dung riêng tư, token đăng nhập hoặc ảnh của người dùng.",
)

add_heading(doc, "9. Kết luận")
add_body(
    doc,
    "Tuần 12 giúp dự án tiến thêm một bước từ mức hoàn thành chức năng sang mức có thể chạy thử. Tính năng AI caption đã có cách xử lý lỗi rõ ràng, dễ kiểm thử và không làm gián đoạn chức năng đăng bài. Việc bổ sung E2E và tối ưu bundle cũng làm hệ thống ổn định hơn khi mở rộng.",
)
add_body(
    doc,
    "Kết quả quan trọng nhất không chỉ là tạo được caption, mà là giữ được quyền quyết định cho người dùng và có phương án khi dịch vụ AI không hoạt động. Đây cũng là cơ sở để tiếp tục theo dõi chất lượng thay vì phụ thuộc hoàn toàn vào kết quả của model.",
)

add_heading(doc, "10. Kế hoạch cho tuần tiếp theo")
add_body(
    doc,
    "Trong tuần 13, em dự kiến đưa hệ thống lên môi trường chạy thử, bổ sung logging và thống kê cho luồng AI, nhưng không lưu nội dung riêng tư. Em sẽ theo dõi thời gian phản hồi, tỷ lệ lỗi, quota và mức độ người dùng chấp nhận caption. Những phản hồi thực tế sẽ được dùng để sửa prompt và giao diện.",
)
add_body(
    doc,
    "Bên cạnh đó, em sẽ mở rộng E2E cho các trường hợp mất mạng, phiên đăng nhập hết hạn và thao tác đồng thời. Phần kiểm tra trên nhiều trình duyệt, quyền camera/microphone và responsive cũng sẽ được thực hiện kỹ hơn trước khi triển khai rộng.",
)

add_heading(doc, "11. Tự đánh giá kết quả tuần")
add_body(
    doc,
    "Khối lượng công việc trong tuần đã đạt mục tiêu đề ra. Phần AI không chỉ chạy được ở trường hợp bình thường mà đã có xử lý lỗi, giới hạn sử dụng và test cho các thành phần chính. Việc giảm bundle và chạy lại các luồng quan trọng giúp em yên tâm hơn khi chuyển sang giai đoạn chạy thử.",
)
add_body(
    doc,
    "Điểm còn thiếu là số lượng tình huống đánh giá caption chưa nhiều và kiểm thử E2E vẫn cần mở rộng trên các môi trường khác nhau. Đây sẽ là nội dung cần ưu tiên trong tuần tiếp theo.",
)

doc.core_properties.title = "Báo cáo tuần 12 - Social Networking Promax"
doc.core_properties.subject = "Hoàn thiện AI caption, kiểm thử E2E và tối ưu hiệu năng"
doc.core_properties.keywords = "Social Networking Promax, tuần 12, AI caption, E2E, hiệu năng"
doc.save(OUTPUT)
print(OUTPUT)
