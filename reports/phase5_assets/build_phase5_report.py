from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "phase5_assets"
OUTPUT = ROOT / "bao_cao_phase_5_realtime_chat.docx"
FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")

COLORS = {
    "ink": "#17324D",
    "muted": "#5D6B7A",
    "blue": "#2E74B5",
    "blue_soft": "#EAF2FA",
    "green": "#2E7D57",
    "green_soft": "#EAF6EF",
    "gold": "#A66A00",
    "gold_soft": "#FFF4DF",
    "red": "#A23B3B",
    "red_soft": "#FFF0F0",
    "gray": "#D8DEE6",
    "gray_soft": "#F5F7FA",
    "white": "#FFFFFF",
    "black": "#000000",
}


def hex_to_rgb(value):
    value = value.lstrip("#")
    return tuple(int(value[i:i + 2], 16) for i in (0, 2, 4))


def pil_font(size, bold=False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(path), size)


def wrap_text(draw, text, font_obj, max_width):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        probe = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), probe, font=font_obj)[2] <= max_width:
            current = probe
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def centered_text(draw, box, text, font_obj, color=COLORS["ink"], max_width=None, line_gap=7):
    x1, y1, x2, y2 = box
    lines = []
    for raw in str(text).split("\n"):
        lines.extend(wrap_text(draw, raw, font_obj, max_width or x2 - x1 - 30) or [""])
    heights = [draw.textbbox((0, 0), line, font=font_obj)[3] - draw.textbbox((0, 0), line, font=font_obj)[1] for line in lines]
    total = sum(heights) + max(0, len(lines) - 1) * line_gap
    y = y1 + (y2 - y1 - total) / 2
    for line, height in zip(lines, heights):
        width = draw.textbbox((0, 0), line, font=font_obj)[2]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font_obj, fill=hex_to_rgb(color))
        y += height + line_gap


def draw_box(draw, box, title, body="", fill=COLORS["blue_soft"], outline=COLORS["blue"], radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=3)
    x1, y1, x2, y2 = box
    title_font = pil_font(26, True)
    body_font = pil_font(21)
    if body:
        draw.text((x1 + 22, y1 + 16), title, font=title_font, fill=hex_to_rgb(COLORS["ink"]))
        y = y1 + 56
        for line in wrap_text(draw, body, body_font, x2 - x1 - 44):
            draw.text((x1 + 22, y), line, font=body_font, fill=hex_to_rgb(COLORS["muted"]))
            y += 29
    else:
        centered_text(draw, box, title, title_font, max_width=x2 - x1 - 28)


def arrow(draw, start, end, color=COLORS["ink"], width=4):
    draw.line([start, end], fill=hex_to_rgb(color), width=width)
    sx, sy = start
    ex, ey = end
    dx = ex - sx
    dy = ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 18
    p1 = (ex, ey)
    p2 = (ex - ux * size + px * size * 0.55, ey - uy * size + py * size * 0.55)
    p3 = (ex - ux * size - px * size * 0.55, ey - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=hex_to_rgb(color))


def label(draw, xy, text, size=21, color=COLORS["muted"]):
    draw.text(xy, str(text), font=pil_font(size), fill=hex_to_rgb(color))


def label_box(draw, xy, text, size=21, color=COLORS["muted"]):
    x, y = xy
    fnt = pil_font(size)
    bbox = draw.textbbox((x, y), str(text), font=fnt)
    pad = 5
    draw.rounded_rectangle(
        (bbox[0] - pad, bbox[1] - pad, bbox[2] + pad, bbox[3] + pad),
        radius=6,
        fill=hex_to_rgb(COLORS["white"]),
    )
    draw.text((x, y), str(text), font=fnt, fill=hex_to_rgb(color))


def make_component_diagram(path):
    img = Image.new("RGB", (1900, 1160), hex_to_rgb(COLORS["white"]))
    draw = ImageDraw.Draw(img)
    draw.text((60, 46), "Biểu đồ component Phase 5 - Realtime Chat", font=pil_font(40, True), fill=hex_to_rgb(COLORS["ink"]))
    label(draw, (60, 102), "REST API dùng để bootstrap dữ liệu; Socket.IO dùng để đồng bộ realtime giữa các user.", 24)

    draw_box(draw, (90, 205, 445, 390), "React UI", "MessengerPage, ConversationList, ChatWindow, MessageBubble, Profile CTA", COLORS["blue_soft"])
    draw_box(draw, (90, 500, 445, 675), "Socket Client", "socket.io-client, emitChatEvent, auto reconnect, ACK timeout", COLORS["green_soft"], COLORS["green"])
    draw_box(draw, (640, 170, 1010, 340), "REST API", "conversation routes, message fallback, unread summary", COLORS["gray_soft"])
    draw_box(draw, (640, 455, 1010, 650), "Socket.IO Server", "JWT middleware, user rooms, conversation rooms, online events", COLORS["green_soft"], COLORS["green"])
    draw_box(draw, (1195, 170, 1565, 365), "Chat Service", "permission, create/reuse conversation, send/reply/delete/read", COLORS["gold_soft"], COLORS["gold"])
    draw_box(draw, (1195, 500, 1565, 675), "Online Users Map", "memory map userId -> socketIds, online/offline friends", COLORS["gray_soft"], COLORS["muted"])
    draw_box(draw, (1195, 805, 1565, 980), "PostgreSQL", "Conversations, ConversationMembers, Messages, Users, Friendships", COLORS["blue_soft"])

    arrow(draw, (445, 275), (640, 250), COLORS["blue"])
    label_box(draw, (495, 195), "REST bootstrap", 20)
    arrow(draw, (445, 590), (640, 545), COLORS["green"])
    label_box(draw, (498, 505), "socket events", 20)
    arrow(draw, (1010, 255), (1195, 260), COLORS["blue"])
    arrow(draw, (1010, 550), (1195, 270), COLORS["green"])
    label_box(draw, (1030, 404), "domain logic", 20)
    arrow(draw, (1380, 365), (1380, 805), COLORS["gold"])
    label_box(draw, (1410, 700), "Sequelize transaction/query", 20)
    arrow(draw, (1010, 555), (1195, 575), COLORS["green"])
    arrow(draw, (1378, 500), (1378, 365), COLORS["muted"])
    label_box(draw, (1410, 430), "online state", 20)
    arrow(draw, (640, 612), (445, 625), COLORS["green"])
    label_box(draw, (496, 646), "broadcast rooms", 20)

    draw.rounded_rectangle((110, 1015, 1780, 1085), radius=16, fill=hex_to_rgb("#FBFCFE"), outline=hex_to_rgb(COLORS["gray"]), width=2)
    label(draw, (140, 1036), "Kết luận kiến trúc: backend là source of truth cho membership, permission, read state và deleted state.", 24, COLORS["ink"])
    img.save(path, "PNG", quality=95)


def draw_actor(draw, x, y, title):
    pen = hex_to_rgb(COLORS["ink"])
    draw.ellipse((x - 23, y, x + 23, y + 46), outline=pen, width=4)
    draw.line((x, y + 46, x, y + 126), fill=pen, width=4)
    draw.line((x - 50, y + 78, x + 50, y + 78), fill=pen, width=4)
    draw.line((x, y + 126, x - 42, y + 188), fill=pen, width=4)
    draw.line((x, y + 126, x + 42, y + 188), fill=pen, width=4)
    centered_text(draw, (x - 145, y + 202, x + 145, y + 270), title, pil_font(24, True), max_width=280)


def use_case_ellipse(draw, box, text, fill=COLORS["blue_soft"]):
    draw.ellipse(box, fill=hex_to_rgb(fill), outline=hex_to_rgb(COLORS["blue"]), width=3)
    centered_text(draw, box, text, pil_font(23, True), max_width=285, line_gap=4)


def make_use_case_diagram(path):
    img = Image.new("RGB", (1900, 1220), hex_to_rgb(COLORS["white"]))
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "Biểu đồ use case Phase 5", font=pil_font(40, True), fill=hex_to_rgb(COLORS["ink"]))
    label(draw, (60, 100), "Tác nhân chính là người dùng đã đăng nhập; hệ thống kiểm tra quan hệ bạn bè và membership trước mỗi thao tác.", 24)

    system = (370, 165, 1505, 1080)
    draw.rounded_rectangle(system, radius=28, outline=hex_to_rgb(COLORS["gray"]), width=4, fill=hex_to_rgb("#FBFCFE"))
    centered_text(draw, (695, 176, 1180, 230), "Social Networking Promax - Chat", pil_font(30, True))

    cases = [
        ((485, 270, 790, 365), "Xem danh sách\nconversation", COLORS["blue_soft"]),
        ((485, 405, 790, 500), "Tạo/reuse\nprivate chat", COLORS["green_soft"]),
        ((485, 540, 790, 635), "Tạo group chat", COLORS["green_soft"]),
        ((485, 675, 790, 770), "Xem lịch sử\ntin nhắn", COLORS["blue_soft"]),
        ((485, 810, 790, 905), "Đánh dấu đã đọc\nvà unread count", COLORS["gold_soft"]),
        ((1040, 270, 1345, 365), "Gửi tin nhắn\nrealtime", COLORS["green_soft"]),
        ((1040, 405, 1345, 500), "Reply tin nhắn", COLORS["green_soft"]),
        ((1040, 540, 1345, 635), "Thu hồi tin nhắn", COLORS["red_soft"]),
        ((1040, 675, 1345, 770), "Theo dõi\nonline/offline", COLORS["gold_soft"]),
        ((1040, 810, 1345, 905), "Mở chat từ\nprofile bạn bè", COLORS["blue_soft"]),
    ]
    def association(points):
        draw.line(points, fill=hex_to_rgb(COLORS["muted"]), width=3, joint="curve")

    for y in [318, 453, 588, 723, 858]:
        association([(280, 570), (340, 570), (340, y), (485, y)])
    for y in [318, 453, 588, 723, 858]:
        association([(280, 570), (340, 570), (340, y), (1040, y)])
    for y in [318, 453, 723, 858]:
        association([(1620, 620), (1565, 620), (1565, y), (1345, y)])

    for box, text, fill in cases:
        use_case_ellipse(draw, box, text, fill)

    draw_actor(draw, 180, 470, "Người dùng")
    draw_actor(draw, 1720, 520, "Bạn bè")

    label(draw, (430, 1015), "Ghi chú: các use case chat riêng tư chỉ hợp lệ khi hai user là accepted friends; non-member không được đọc/gửi/xóa tin nhắn.", 22, COLORS["ink"])
    img.save(path, "PNG", quality=95)


def make_erd_diagram(path):
    img = Image.new("RGB", (1900, 1120), hex_to_rgb(COLORS["white"]))
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "Biểu đồ dữ liệu Phase 5", font=pil_font(40, True), fill=hex_to_rgb(COLORS["ink"]))
    label(draw, (60, 100), "Ba bảng mới được thêm cho chat và liên kết với Users/Friendships đã có từ Phase 1-2.", 24)

    def entity(box, title, fields, fill=COLORS["blue_soft"], outline=COLORS["blue"]):
        draw.rounded_rectangle(box, radius=18, fill=hex_to_rgb(fill), outline=hex_to_rgb(outline), width=3)
        x1, y1, x2, y2 = box
        draw.rectangle((x1, y1, x2, y1 + 54), fill=hex_to_rgb(outline))
        centered_text(draw, (x1, y1 + 5, x2, y1 + 50), title, pil_font(24, True), COLORS["white"], max_width=x2 - x1 - 24)
        y = y1 + 75
        for field in fields:
            draw.text((x1 + 22, y), field, font=pil_font(20), fill=hex_to_rgb(COLORS["ink"]))
            y += 30

    entity((115, 240, 500, 515), "Users", ["id PK", "username, fullName", "avatar, bio", "createdAt"], COLORS["gray_soft"], COLORS["muted"])
    entity((695, 190, 1115, 515), "Conversations", ["id PK", "type: private/group", "name, avatar", "adminId FK Users", "lastMessageAt"], COLORS["blue_soft"], COLORS["blue"])
    entity((1320, 190, 1755, 575), "ConversationMembers", ["id PK", "conversationId FK", "userId FK", "role: member/admin", "lastReadAt", "deletedAt", "UNIQUE(conversationId,userId)"], COLORS["green_soft"], COLORS["green"])
    entity((695, 680, 1115, 1010), "Messages", ["id PK", "conversationId FK", "senderId FK Users", "replyToId FK Messages", "content, media", "type, isDeleted, deletedAt"], COLORS["gold_soft"], COLORS["gold"])
    entity((115, 680, 500, 900), "Friendships", ["requesterId FK Users", "addresseeId FK Users", "status = accepted", "dùng để chặn spam chat"], COLORS["gray_soft"], COLORS["muted"])

    arrow(draw, (500, 377), (695, 345), COLORS["muted"])
    label_box(draw, (555, 397), "FK Users", 20)
    arrow(draw, (1115, 350), (1320, 350), COLORS["green"])
    label_box(draw, (1162, 310), "1 - N members", 20)
    arrow(draw, (905, 515), (905, 680), COLORS["gold"])
    label_box(draw, (925, 580), "1 - N messages", 20)
    arrow(draw, (500, 790), (695, 845), COLORS["muted"])
    label_box(draw, (535, 875), "accepted friends", 20)
    arrow(draw, (1320, 438), (1115, 825), COLORS["green"])
    label_box(draw, (1175, 628), "membership gates read/write", 20)

    img.save(path, "PNG", quality=95)


def make_sequence_diagram(path):
    img = Image.new("RGB", (2000, 1250), hex_to_rgb(COLORS["white"]))
    draw = ImageDraw.Draw(img)
    draw.text((60, 44), "Biểu đồ sequence: gửi tin nhắn realtime", font=pil_font(40, True), fill=hex_to_rgb(COLORS["ink"]))
    label(draw, (60, 100), "Luồng thể hiện cách tin nhắn được validate, lưu DB, ACK về sender và broadcast tới members trong conversation.", 24)

    participants = [
        ("Sender UI", 170),
        ("Socket Client", 440),
        ("Socket.IO Server", 750),
        ("Chat Service", 1075),
        ("PostgreSQL", 1390),
        ("Receiver UI", 1710),
    ]
    top = 190
    bottom = 1125
    for name, x in participants:
        draw.rounded_rectangle((x - 105, top, x + 105, top + 62), radius=14, fill=hex_to_rgb(COLORS["blue_soft"]), outline=hex_to_rgb(COLORS["blue"]), width=3)
        centered_text(draw, (x - 100, top + 6, x + 100, top + 56), name, pil_font(21, True), max_width=185)
        draw.line((x, top + 62, x, bottom), fill=hex_to_rgb(COLORS["gray"]), width=3)

    def seq(y, start_x, end_x, text, color=COLORS["ink"], dashed=False):
        if dashed:
            direction = 1 if end_x > start_x else -1
            x = start_x
            while (end_x - x) * direction > 0:
                nx = x + direction * min(24, abs(end_x - x))
                draw.line((x, y, nx, y), fill=hex_to_rgb(color), width=3)
                x = nx + direction * 12
            arrow(draw, (end_x - direction * 1, y), (end_x, y), color, 3)
        else:
            arrow(draw, (start_x, y), (end_x, y), color, 3)
        mid = min(start_x, end_x) + abs(end_x - start_x) / 2
        for offset, line in enumerate(wrap_text(draw, text, pil_font(19), 260)):
            draw.text((mid - 130, y - 48 + offset * 24), line, font=pil_font(19), fill=hex_to_rgb(COLORS["ink"]))

    seq(325, 170, 440, "1. Nhập content, chọn reply target nếu có")
    seq(415, 440, 750, "2. emit send_message/reply_message + ACK")
    seq(505, 750, 1075, "3. validate payload, ensureMember()")
    seq(595, 1075, 1390, "4. transaction: create Message, update lastMessageAt, lastReadAt")
    seq(685, 1390, 1075, "5. trả message kèm sender/replyTo", COLORS["green"], True)
    seq(775, 1075, 750, "6. serialize public message", COLORS["green"], True)
    seq(865, 750, 440, "7. ACK ok: true, message", COLORS["green"], True)
    seq(955, 750, 1710, "8. emit new_message/new_reply tới conversation room và user rooms", COLORS["green"], True)
    seq(1045, 1710, 1710, "9. UI append bubble, update unread/lastMessage", COLORS["gold"])
    img.save(path, "PNG", quality=95)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths):
    set_table_width(table, widths)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    set_run_font(run, size=10.5, color=COLORS["black"])
            if i == 0:
                set_cell_shading(cell, COLORS["gray_soft"])
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_para(doc, text="", style=None, bold=False, italic=False, color=None, size=None, align=None, after=None, before=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    return p


def add_bullets(doc, items):
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        add_para(doc, item, style="List Number")


def add_caption(doc, text):
    p = add_para(doc, text, style="Caption", italic=True, color="#555555", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    return p


def add_picture(doc, path, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    inline._inline.docPr.set("descr", path.stem.replace("_", " "))
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    style_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def apply_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, "#1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167

    if "Caption" in styles:
        styles["Caption"].font.name = "Calibri"
        styles["Caption"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        styles["Caption"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        styles["Caption"].font.size = Pt(10)
        styles["Caption"].font.italic = True
        styles["Caption"].font.color.rgb = RGBColor(85, 85, 85)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Báo cáo Phase 5 - Social Networking Promax")
    set_run_font(r, size=9, color=COLORS["muted"])


def add_cover(doc):
    add_para(doc, "BÁO CÁO TIẾN ĐỘ DỰ ÁN", bold=True, size=24, color=COLORS["ink"], align=WD_ALIGN_PARAGRAPH.CENTER, before=80, after=10)
    add_para(doc, "Social Networking Promax", bold=True, size=20, color=COLORS["blue"], align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Phase 5 - Realtime Chat", bold=True, size=16, color=COLORS["muted"], align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
    add_para(doc, "Nội dung báo cáo: Xây dựng chức năng nhắn tin realtime, quản lý conversation, message, Socket.IO và giao diện Messenger.", size=12, color=COLORS["ink"], align=WD_ALIGN_PARAGRAPH.CENTER, after=42)
    table = add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ["Sinh viên", "Đường Gia Bằng"],
            ["MSV", "22026537"],
            ["Giai đoạn", "Phase 5 - Realtime Chat"],
            ["Thời gian trong kế hoạch", "Tuần 7-8"],
            ["Dự án", "Social Networking Promax"],
        ],
        [1.8, 4.5],
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_para(doc, "Hà Nội, 2026", bold=True, size=12, color=COLORS["ink"], align=WD_ALIGN_PARAGRAPH.CENTER, before=45)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_footer(doc.sections[-1])


def build_report():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    component = ASSET_DIR / "phase5_component_diagram.png"
    use_case = ASSET_DIR / "phase5_use_case_diagram.png"
    erd = ASSET_DIR / "phase5_erd_diagram.png"
    sequence = ASSET_DIR / "phase5_sequence_diagram.png"
    make_component_diagram(component)
    make_use_case_diagram(use_case)
    make_erd_diagram(erd)
    make_sequence_diagram(sequence)

    doc = Document()
    apply_styles(doc)
    add_footer(doc.sections[0])
    add_cover(doc)

    add_para(doc, "1. Tổng quan Phase 5", style="Heading 1")
    add_para(
        doc,
        "Phase 5 của dự án Social Networking Promax tập trung xây dựng chức năng nhắn tin realtime. "
        "Sau khi các phase trước đã hoàn thiện nền tảng xác thực, hồ sơ cá nhân, quan hệ bạn bè, bài viết và story, Phase 5 bổ sung kênh giao tiếp trực tiếp giữa người dùng. "
        "Đây là một phần quan trọng của mạng xã hội vì giúp user có thể trao đổi riêng tư, theo dõi trạng thái online và nhận tin nhắn mới mà không cần tải lại trang."
    )
    add_para(
        doc,
        "Về mặt kỹ thuật, Phase 5 không chỉ thêm giao diện Messenger mà còn đưa vào một lớp realtime mới dựa trên Socket.IO. "
        "REST API vẫn được dùng cho các thao tác lấy dữ liệu ban đầu, phân trang lịch sử tin nhắn và fallback khi socket không sẵn sàng. "
        "Socket.IO đảm nhiệm các sự kiện realtime như gửi tin nhắn, reply, thu hồi tin nhắn, đánh dấu đã đọc và cập nhật online/offline."
    )

    add_para(doc, "2. Mục tiêu, phạm vi và kết quả chính", style="Heading 1")
    add_table(
        doc,
        ["Nhóm mục tiêu", "Nội dung thực hiện"],
        [
            ["Conversation", "Tạo và lấy danh sách conversation; hỗ trợ private chat reuse; hỗ trợ group conversation có admin và members."],
            ["Message", "Lấy lịch sử tin nhắn có phân trang; gửi tin nhắn text; reply message; thu hồi tin nhắn bằng soft delete."],
            ["Realtime", "Thiết lập Socket.IO server/client; xác thực socket bằng JWT; broadcast message theo user room và conversation room."],
            ["Read/Unread", "Cập nhật lastReadAt, tính unreadCount từng conversation và tổng unread count trên navbar."],
            ["Online", "Theo dõi user online/offline bằng memory map và broadcast trạng thái cho bạn bè accepted."],
            ["Frontend", "Xây dựng trang /messenger gồm conversation list, chat window, message bubble, reply bar, load older, empty/loading/error state và responsive layout."],
            ["Kiểm thử", "Backend chat service test pass 9/9; frontend production build pass."],
        ],
        [1.65, 4.75],
    )
    add_para(
        doc,
        "Một số phần được xác định là ngoài phạm vi hoặc để phát triển tiếp: chat media/file, notification push/FCM, search message, message reaction, pinned message và typing indicator hoàn chỉnh. "
        "Những phần này có thể được mở rộng sau khi core chat ổn định."
    )

    add_para(doc, "3. Kiến trúc realtime chat", style="Heading 1")
    add_para(
        doc,
        "Kiến trúc Phase 5 được tách thành hai luồng chính. Luồng REST dùng cho các thao tác cần tính ổn định và dễ kiểm thử như lấy danh sách conversation, lấy message history, tạo conversation, lấy unread summary và fallback gửi/xóa tin. "
        "Luồng Socket.IO dùng cho các thao tác cần cập nhật ngay lập tức như gửi tin, reply, thu hồi tin, mark read và online/offline. Cách tách này giúp hệ thống vừa giữ được tính realtime, vừa không phụ thuộc hoàn toàn vào kết nối websocket."
    )
    add_picture(doc, component, 6.4)
    add_caption(doc, "Hình 1. Biểu đồ component tổng quát của Phase 5.")
    add_table(
        doc,
        ["Component", "Vai trò trong Phase 5"],
        [
            ["React UI", "Cung cấp trải nghiệm Messenger: danh sách chat, cửa sổ chat, bubble tin nhắn, trạng thái online và điều hướng từ profile bạn bè."],
            ["Socket client", "Kết nối Socket.IO bằng access token, gửi event kèm ACK, timeout khi server không phản hồi."],
            ["REST API", "Bootstrap dữ liệu conversation/message, cung cấp unread-count và fallback khi socket không hoạt động."],
            ["Socket.IO server", "Xác thực JWT, quản lý room, nhận event từ client, gọi service và broadcast kết quả cho các member."],
            ["Chat service", "Tập trung logic nghiệp vụ: kiểm tra quyền, tạo/reuse conversation, gửi/reply/delete/read message."],
            ["PostgreSQL", "Lưu dữ liệu conversation, member, message và liên kết với Users/Friendships."],
        ],
        [1.65, 4.75],
    )

    add_para(doc, "4. Thiết kế dữ liệu", style="Heading 1")
    add_para(
        doc,
        "Phase 5 bổ sung ba bảng chính: Conversations, ConversationMembers và Messages. "
        "Thiết kế này giúp tách rõ thực thể cuộc trò chuyện, thành viên tham gia và nội dung tin nhắn. "
        "ConversationMembers đóng vai trò rất quan trọng vì mọi quyền đọc, gửi, xóa, mark read đều phải đi qua membership."
    )
    add_picture(doc, erd, 6.4)
    add_caption(doc, "Hình 2. Biểu đồ dữ liệu cho chức năng realtime chat.")
    add_table(
        doc,
        ["Bảng", "Các trường chính", "Ý nghĩa"],
        [
            ["Conversations", "id, type, name, avatar, adminId, lastMessageAt", "Lưu thông tin conversation; lastMessageAt dùng để sort danh sách chat mới nhất."],
            ["ConversationMembers", "conversationId, userId, role, joinedAt, lastReadAt, deletedAt", "Lưu membership, role admin/member và read state của từng user trong conversation."],
            ["Messages", "conversationId, senderId, replyToId, content, media, type, isDeleted, deletedAt", "Lưu tin nhắn, reply relation, trạng thái thu hồi và dữ liệu media mở rộng cho phase sau."],
            ["Friendships", "requesterId, addresseeId, status", "Dùng để kiểm tra chỉ accepted friends mới tạo private/group chat trong phạm vi Phase 5."],
        ],
        [1.35, 2.25, 2.8],
    )
    add_para(
        doc,
        "Migration cũng tạo các index quan trọng trên type, adminId, lastMessageAt, conversationId, userId, senderId, replyToId và cặp conversationId-createdAt. "
        "Riêng bảng ConversationMembers có unique constraint trên conversationId + userId để tránh một user bị thêm trùng vào cùng một conversation."
    )

    add_para(doc, "5. API và Socket.IO", style="Heading 1")
    add_para(doc, "5.1. REST API", style="Heading 2")
    add_table(
        doc,
        ["Endpoint", "Mục đích", "Ghi chú bảo mật/nghiệp vụ"],
        [
            ["GET /api/conversations", "Lấy danh sách conversation của current user.", "Chỉ trả conversation mà user là member; sort theo lastMessageAt."],
            ["GET /api/conversations/unread-count", "Lấy tổng số tin chưa đọc.", "Tính từ lastReadAt của từng ConversationMember."],
            ["POST /api/conversations", "Tạo private/group conversation.", "Private chat reuse nếu đã tồn tại; chỉ cho accepted friends."],
            ["GET /api/conversations/:id/messages", "Lấy message history có cursor/limit.", "Check membership trước khi đọc; response trả oldest-first cho UI render tự nhiên."],
            ["POST /api/conversations/:id/messages", "Fallback gửi tin nhắn qua REST.", "Validate content 1-5000 ký tự và replyToId nếu có."],
            ["DELETE /api/messages/:id", "Thu hồi tin nhắn.", "Chỉ sender được xóa; dùng soft delete, không xóa row khỏi DB."],
            ["GET /api/users/online-friends", "Lấy danh sách bạn bè đang online.", "Dựa trên memory map của Socket.IO server."],
        ],
        [2.05, 2.25, 2.1],
    )

    add_para(doc, "5.2. Socket.IO events", style="Heading 2")
    add_table(
        doc,
        ["Event", "Chiều", "Nội dung xử lý"],
        [
            ["connect", "Client -> Server", "Client gửi access token trong auth; server verify JWT và join room user:{id}."],
            ["join_conversation", "Client -> Server", "Server ensureMember rồi cho socket join room conversation:{conversationId}."],
            ["send_message", "Client -> Server", "Tạo message text, cập nhật lastMessageAt/lastReadAt và broadcast new_message."],
            ["reply_message", "Client -> Server", "Tạo message có replyToId và broadcast new_reply."],
            ["delete_message", "Client -> Server", "Soft delete message của sender và broadcast message_deleted."],
            ["mark_read", "Client -> Server", "Cập nhật lastReadAt và broadcast message_read."],
            ["user_online/user_offline", "Server -> Client", "Broadcast trạng thái online/offline tới accepted friends."],
            ["online_users_list", "Server -> Client", "Gửi danh sách friend đang online ngay sau khi socket connect."],
            ["socket_error", "Server -> Client", "Trả lỗi domain như CONVERSATION_FORBIDDEN, MESSAGE_DELETE_FORBIDDEN."],
        ],
        [1.7, 1.45, 3.25],
    )
    add_para(
        doc,
        "Socket event có ACK đối với các thao tác gửi/reply/xóa message. Frontend dùng helper emitChatEvent với timeout 7 giây để tránh treo UI nếu server không phản hồi. "
        "Nếu socket không connected, frontend có fallback REST cho gửi và xóa tin nhắn."
    )

    add_para(doc, "6. Các chức năng đã hoàn thành", style="Heading 1")
    add_para(doc, "6.1. Conversation", style="Heading 2")
    add_bullets(doc, [
        "User đăng nhập có thể lấy danh sách conversation của mình, bao gồm members, lastMessage, unreadCount, avatar/name hiển thị và lastMessageAt.",
        "Private conversation được reuse nếu đã tồn tại giữa hai user, giúp tránh tạo nhiều box chat trùng.",
        "Chỉ accepted friends mới được tạo private chat hoặc thêm vào group chat, giảm spam và bám sát social graph của hệ thống.",
        "Group conversation yêu cầu ít nhất hai member khác ngoài creator; creator được đặt role admin.",
    ])
    add_para(doc, "6.2. Message", style="Heading 2")
    add_bullets(doc, [
        "Message history có phân trang bằng cursor và limit, query newest-first nhưng trả về oldest-first để UI hiển thị đúng thứ tự đọc.",
        "Gửi tin nhắn text qua Socket.IO hoặc REST fallback, validate content từ 1 đến 5000 ký tự.",
        "Reply message được lưu bằng replyToId và trả kèm thông tin sender của message gốc để UI render reply preview.",
        "Thu hồi tin nhắn dùng soft delete: content/media bị null, isDeleted=true, deletedAt có giá trị; UI hiển thị trạng thái 'Tin nhắn đã được thu hồi'.",
    ])
    add_para(doc, "6.3. Read, unread và online", style="Heading 2")
    add_bullets(doc, [
        "Khi user gửi hoặc mở conversation, hệ thống cập nhật lastReadAt để tính unreadCount chính xác hơn.",
        "Unread count được tính theo số message của người khác có createdAt lớn hơn lastReadAt.",
        "Navbar có tổng unread badge thông qua Redux slice chatNotifications.",
        "Online state được quản lý bằng memory map userId -> socketIds, hỗ trợ user mở nhiều tab/socket cùng lúc.",
        "Khi user online/offline, server chỉ broadcast cho accepted friends thay vì toàn bộ hệ thống.",
    ])
    add_para(doc, "6.4. Giao diện Messenger", style="Heading 2")
    add_bullets(doc, [
        "Trang /messenger dùng layout hai cột trên desktop: conversation list bên trái và chat window bên phải.",
        "Mobile chuyển về layout một cột; khi chọn conversation thì ẩn list và có nút quay lại.",
        "Conversation list hiển thị avatar, tên conversation, preview last message, unread badge và online dot.",
        "Chat window có header, avatar, trạng thái online/offline, nút video call cho phase sau, message list, load older và input gửi tin.",
        "Message bubble phân biệt tin của mình/tin của người khác, hỗ trợ copy, reply và delete với hover actions.",
        "Từ profile bạn bè có nút Nhắn tin để tạo/reuse private conversation rồi điều hướng sang /messenger?conversationId=...",
    ])

    add_para(doc, "7. Biểu đồ use case", style="Heading 1")
    add_para(
        doc,
        "Biểu đồ use case dưới đây mô tả các thao tác chính của người dùng trong Phase 5. "
        "Tất cả thao tác chat đều yêu cầu user đã đăng nhập. Một số thao tác như tạo private chat hoặc mở chat từ profile còn phụ thuộc điều kiện hai user đã là bạn bè accepted."
    )
    add_picture(doc, use_case, 6.4)
    add_caption(doc, "Hình 3. Biểu đồ use case của chức năng realtime chat.")

    add_para(doc, "8. Biểu đồ sequence", style="Heading 1")
    add_para(
        doc,
        "Luồng gửi tin nhắn realtime là luồng quan trọng nhất của Phase 5. "
        "Client gửi event qua Socket.IO, server kiểm tra quyền bằng ensureMember, service lưu message trong database, sau đó server ACK về sender và broadcast message mới tới các member trong conversation."
    )
    add_picture(doc, sequence, 6.4)
    add_caption(doc, "Hình 4. Biểu đồ sequence cho luồng gửi tin nhắn realtime.")
    add_numbered(doc, [
        "Người dùng nhập nội dung tin nhắn ở ChatWindow và bấm gửi.",
        "Frontend gọi emitChatEvent với send_message hoặc reply_message, kèm conversationId, content và replyToId nếu có.",
        "Socket.IO server nhận event, lấy user từ socket.data.user và gọi chatService.createMessage.",
        "Service validate content, kiểm tra user là member, kiểm tra replyToId thuộc cùng conversation nếu có.",
        "Service tạo message trong transaction, cập nhật lastMessageAt và lastReadAt.",
        "Server broadcast new_message hoặc new_reply tới conversation room và user room của các member.",
        "Frontend cập nhật message list, lastMessage, sort conversation lên đầu và cập nhật unread nếu conversation không active.",
    ])

    add_para(doc, "9. Kiểm thử và chất lượng", style="Heading 1")
    add_para(
        doc,
        "Phase 5 có kiểm thử backend tập trung vào chat service vì đây là nơi chứa logic quan trọng nhất: permission, private reuse, group create, send/delete/read và unread summary. "
        "Ngoài ra frontend đã được build production để kiểm tra code UI Messenger có thể đóng gói thành công."
    )
    add_table(
        doc,
        ["Hạng mục", "Kết quả", "Ý nghĩa"],
        [
            ["Backend test", "npm test -- chat.service.test.js: PASS 9/9 tests", "Xác nhận logic conversation/message/read/unread hoạt động theo testcase service."],
            ["Frontend build", "npm run build: PASS", "Xác nhận React/Vite build được production bundle. Có warning chunk > 500 kB, chưa phải lỗi build."],
            ["Validation", "express-validator + service validation", "Chặn payload sai như UUID không hợp lệ, content rỗng/quá dài, group thiếu member/name."],
            ["Permission", "ensureMember + sender check", "Non-member không đọc/gửi; non-sender không được thu hồi message."],
            ["Regression scope", "Auth, friends, navbar, profile CTA", "Phase 5 phụ thuộc Phase 1-2 nên cần đảm bảo auth/friends vẫn ổn."],
        ],
        [1.45, 2.1, 2.85],
    )
    add_para(
        doc,
        "Kết quả test thực tế đã chạy trong quá trình lập báo cáo: backend chat service pass 9 test, frontend build pass. "
        "Warning về kích thước bundle có thể xử lý sau bằng dynamic import hoặc manualChunks nếu dự án cần tối ưu performance production."
    )

    add_para(doc, "10. Kiến thức học được", style="Heading 1")
    add_para(
        doc,
        "Phase 5 giúp hiểu sâu hơn sự khác biệt giữa ứng dụng CRUD thông thường và ứng dụng realtime. "
        "Khi có nhiều user cùng tương tác trên cùng một conversation, dữ liệu không chỉ được lưu đúng mà còn phải được đồng bộ đúng thời điểm, đúng người nhận và đúng quyền."
    )
    add_table(
        doc,
        ["Nhóm kiến thức", "Nội dung học được"],
        [
            ["Socket.IO", "Biết cách tạo socket server, xác thực kết nối bằng JWT, join room theo user/conversation và broadcast event realtime."],
            ["REST + Realtime", "Hiểu vì sao không thay toàn bộ REST bằng socket: REST phù hợp bootstrap/pagination/fallback, socket phù hợp đồng bộ sự kiện tức thời."],
            ["Permission realtime", "Mọi socket event vẫn cần kiểm tra quyền ở backend. Không thể tin client chỉ vì client đã mở đúng màn hình."],
            ["Thiết kế dữ liệu chat", "ConversationMembers là bảng trung tâm để quản lý membership, role, lastReadAt và deletedAt. Messages cần soft delete để giữ lịch sử."],
            ["Unread/read state", "Unread count không nên chỉ lưu cứng ở frontend; backend cần tính theo lastReadAt và message của người khác."],
            ["Online state", "Memory map dễ triển khai cho single server, nhưng khi scale multi-instance cần Redis adapter hoặc shared presence store."],
            ["UI state sync", "Frontend cần tránh duplicate message khi vừa nhận ACK vừa nhận broadcast, đồng thời phải sort conversation sau message mới."],
            ["Testing service layer", "Việc gom logic vào chat.service giúp test unit dễ hơn so với nhét logic vào controller/socket handler."],
        ],
        [1.55, 4.85],
    )
    add_bullets(doc, [
        "Học được cách thiết kế ACK timeout để UI không bị kẹt khi socket request không phản hồi.",
        "Học được cách dùng room user:{id} để gửi event cá nhân và room conversation:{id} để đồng bộ trong một cuộc trò chuyện.",
        "Học được cách kết hợp optimistic UI vừa phải: chỉ append message khi có response/ACK để tránh hiển thị dữ liệu chưa được server xác nhận.",
        "Học được cách tách phần realtime core trước, còn các phần nâng cao như typing, media, push notification có thể mở rộng sau để tránh phình scope.",
    ])

    add_para(doc, "11. Hạn chế và hướng phát triển", style="Heading 1")
    add_bullets(doc, [
        "Typing indicator mới nằm trong yêu cầu/testcase, chưa thấy handler rõ trong code hiện tại; có thể bổ sung typing và stop_typing ở cả socket server và ChatWindow.",
        "Online users map đang lưu trong memory nên phù hợp single server. Khi deploy nhiều instance cần Redis adapter cho Socket.IO hoặc một presence store dùng chung.",
        "Chat hiện tập trung vào text message. Media/file message có cột media và type chuẩn bị sẵn nhưng cần thêm upload security, preview UI và giới hạn dung lượng.",
        "Chưa có search message, message reaction, pinned message và message encryption.",
        "Frontend build pass nhưng bundle JS lớn hơn 500 kB; có thể tách route Messenger bằng dynamic import để tối ưu tải trang.",
        "Socket integration/E2E test có thể bổ sung thêm để kiểm tra realtime giữa hai client thật.",
    ])

    add_para(doc, "12. Kết luận", style="Heading 1")
    add_para(
        doc,
        "Sau Phase 5, dự án Social Networking Promax đã có nền tảng nhắn tin realtime tương đối hoàn chỉnh. "
        "Hệ thống đã có conversation model, member permission, message history, gửi/reply/thu hồi tin, read/unread state, online/offline state và giao diện Messenger responsive. "
        "Việc tách rõ REST API, Socket.IO server và chat service giúp code dễ mở rộng sang Phase 6 video call và Phase 7 notifications."
    )
    add_para(
        doc,
        "Kết quả Phase 5 cũng tạo ra nhiều kinh nghiệm quan trọng về realtime system: xác thực socket, room design, broadcast đúng phạm vi, xử lý fallback, kiểm soát quyền ở service layer và đồng bộ state frontend. "
        "Các hạn chế còn lại như typing indicator đầy đủ, chat media/file, push notification và scaling online presence có thể tiếp tục phát triển trong các phase sau."
    )

    doc.core_properties.title = "Báo cáo Phase 5 - Realtime Chat"
    doc.core_properties.subject = "Social Networking Promax"
    doc.core_properties.author = "Đường Gia Bằng"
    doc.save(OUTPUT)
    return OUTPUT


def structural_check(path):
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        media = [name for name in names if name.startswith("word/media/")]
    required = [
        "Báo cáo Phase 5",
        "Kiến thức học được",
        "Biểu đồ use case",
        "Biểu đồ sequence",
        "Realtime Chat",
    ]
    missing = [item for item in required if item not in xml]
    if missing:
        raise RuntimeError(f"Missing expected content: {missing}")
    if len(media) < 4:
        raise RuntimeError("Expected at least four embedded diagrams")


if __name__ == "__main__":
    output = build_report()
    structural_check(output)
    print(output)
